import os
import tempfile
from typing import List, Dict, Any, Optional
import docx
from docx.document import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from app.Logger.ocr_logger import get_standard_logger
from app.Exceptions.custom_exceptions import (
    handle_general_operations, 
    log_method_entry_exit, 
    ExceptionSeverity,
    BaseOCRException
)

class DocxExtractionService:
    """
    Service for extracting tables from DOCX files while preserving original formatting.
    This service works alongside the existing PDF/Image processing without modifying it.
    """
    _instance = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super(DocxExtractionService, cls).__new__(cls)
            cls._instance._initialized = False
        return cls._instance
    
    def __init__(self):
        if not self._initialized:
            self.logger = get_standard_logger("DocxExtractionService")
            self._initialized = True
    
    @log_method_entry_exit
    @handle_general_operations(severity=ExceptionSeverity.HIGH)
    def extract_tables_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract tables from a DOCX file while preserving formatting.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of tables with their data and metadata
            
        Raises:
            FileProcessingException: If there's an error processing the file
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"DOCX file not found: {file_path}")
                
            doc = docx.Document(file_path)
            tables = []
            
            # Process all tables in the document
            for table_idx, table in enumerate(doc.tables, 1):
                table_data = self._process_docx_table(table, table_idx)
                tables.append(table_data)
                
            self.logger.info(f"Extracted {len(tables)} tables from DOCX: {file_path}")
            return tables
            
        except Exception as e:
            error_msg = f"Error processing DOCX file {file_path}: {str(e)}"
            self.logger.error(error_msg, exc_info=True)
            raise BaseOCRException(error_msg) from e
    
    def _process_docx_table(self, table: DocxTable, table_index: int) -> Dict[str, Any]:
        """Process a single DOCX table and return its data with formatting."""
        table_data = {
            "table_index": table_index,
            "rows": [],
            "column_count": 0,
            "row_count": 0,
            "formatting": {}
        }
        
        # Get table dimensions
        row_count = len(table.rows)
        col_count = max(len(row.cells) for row in table.rows) if row_count > 0 else 0
        
        table_data["row_count"] = row_count
        table_data["column_count"] = col_count
        
        # Extract cell data
        for row_idx, row in enumerate(table.rows):
            row_data = []
            
            for cell_idx, cell in enumerate(row.cells):
                # Get cell text with formatting
                cell_text = self._get_cell_text(cell)
                
                # Get cell formatting
                cell_format = self._get_cell_format(cell)
                
                row_data.append({
                    "text": cell_text,
                    "format": cell_format,
                    "row_span": 1,  # DOCX handles rowspans/colspans differently
                    "col_span": 1,
                    "is_header": self._is_header_cell(row_idx, cell_idx, table_data)
                })
            
            table_data["rows"].append(row_data)
        
        return table_data
    
    def _get_cell_text(self, cell) -> str:
        """Extract text from a cell, preserving line breaks."""
        return "\n".join(paragraph.text for paragraph in cell.paragraphs if paragraph.text.strip())
    
    def _get_cell_format(self, cell) -> Dict[str, Any]:
        """Extract formatting information from a cell."""
        if not cell.paragraphs:
            return {}
            
        # Get formatting from the first paragraph in the cell
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else None
        
        if not run:
            return {}
            
        return {
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_size": run.font.size.pt if run.font.size else None,
            "font_name": run.font.name,
            "alignment": para.alignment.type if para.alignment else None
        }
    
    def _is_header_cell(self, row_idx: int, cell_idx: int, table_data: Dict) -> bool:
        """Determine if a cell is likely a header cell."""
        # First row is typically a header
        if row_idx == 0:
            return True
            
        # If this cell is in the first column, it might be a row header
        if cell_idx == 0:
            return True
            
        return False

# Singleton instance
docx_extraction_service = DocxExtractionService()
