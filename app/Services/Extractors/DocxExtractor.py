from app.Services.Extractors.BaseDocumentExtractor import BaseDocumentExtractor


class DocxExtractor(BaseDocumentExtractor):
    """DOCX document extractor"""
    
    async def extract_text(self, file_path: str) -> str:
        """Extract all text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as a string
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            all_text_parts = []
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    all_text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        all_text_parts.append(" | ".join(row_texts))
            
            combined_text = "\n".join(all_text_parts)
            
            self.logger.info(f"Extracted {len(combined_text)} characters from DOCX: {file_path}")
            return self.clean_text(combined_text)
            
        except ImportError:
            self.logger.error("python-docx library not available. Install it with: pip install python-docx")
            return "Error: python-docx library not installed"
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX {file_path}: {e}", exc_info=True)
            return f"Error extracting text from DOCX: {str(e)}"

