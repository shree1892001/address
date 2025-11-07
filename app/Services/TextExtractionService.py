import re
import os
import statistics
from typing import List, Dict, Optional
from app.Logger.ocr_logger import get_standard_logger
from app.Exceptions.custom_exceptions import (
    handle_general_operations, log_method_entry_exit, ExceptionSeverity
)
from app.Services.OCRService import OCRService
from app.Services.Extractors.DocumentExtractorFactory import document_extractor_factory


class TextExtractionService:
    _instance = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super(TextExtractionService, cls).__new__(cls)
            cls._instance._initialized = False
        return cls._instance
    
    def __init__(self):
        if not self._initialized:
            self.logger = get_standard_logger("TextExtractionService")
            self.ocr_service = OCRService()
            self._initialized = True
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type based on extension and content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            File type: 'pdf', 'docx', 'jpeg', 'jpg', 'png', or 'unknown'
        """
        if not file_path:
            return 'unknown'
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Map extensions to file types
        extension_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',  # Treat .doc as .docx for extraction
            '.jpeg': 'jpeg',
            '.jpg': 'jpeg',
            '.png': 'png',
            '.gif': 'png',  # Treat GIF as image
            '.bmp': 'png',  # Treat BMP as image
            '.tiff': 'png',  # Treat TIFF as image
            '.tif': 'png'   # Treat TIF as image
        }
        
        file_type = extension_map.get(file_ext, 'unknown')
        
        # Additional validation: check file content for images
        if file_type in ['jpeg', 'png']:
            try:
                with open(file_path, 'rb') as f:
                    # Check image magic numbers
                    header = f.read(8)
                    if file_type == 'jpeg' and not header.startswith(b'\xff\xd8'):
                        self.logger.warning(f"File extension suggests JPEG but magic number doesn't match: {file_path}")
                    elif file_type == 'png' and not header.startswith(b'\x89PNG\r\n\x1a\n'):
                        self.logger.warning(f"File extension suggests PNG but magic number doesn't match: {file_path}")
            except Exception as e:
                self.logger.debug(f"Could not verify image magic number: {e}")
        
        self.logger.info(f"Detected file type: {file_type} for {file_path}")
        return file_type
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract all text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as a string
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            all_text_parts = []
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    all_text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        all_text_parts.append(" | ".join(row_texts))
            
            combined_text = "\n".join(all_text_parts)
            
            self.logger.info(f"Extracted {len(combined_text)} characters from DOCX: {file_path}")
            return combined_text.strip()
            
        except ImportError:
            self.logger.error("python-docx library not available. Install it with: pip install python-docx")
            return "Error: python-docx library not installed"
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX {file_path}: {e}", exc_info=True)
            return f"Error extracting text from DOCX: {str(e)}"
    
    def _extract_text_from_image(self, file_path: str) -> str:
        """Extract text from an image file (JPEG, PNG) using OCR.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            Extracted text as a string
        """
        try:
            from PIL import Image
            import pytesseract
            
            # Open image
            img = Image.open(file_path)
            
            # Perform OCR
            ocr_text = pytesseract.image_to_string(img, lang='eng')
            
            if ocr_text.strip():
                self.logger.info(f"OCR extracted {len(ocr_text)} characters from image: {file_path}")
                return ocr_text.strip()
            else:
                self.logger.warning(f"No text found in image: {file_path}")
                return "No text content found in image"
                
        except ImportError:
            self.logger.error("Required libraries not available. Install with: pip install Pillow pytesseract")
            return "Error: Required OCR libraries not installed"
        except Exception as e:
            self.logger.error(f"Error performing OCR on image {file_path}: {e}", exc_info=True)
            # Try using OCRService as fallback
            try:
                from PIL import Image
                import fitz
                
                # Open image to get dimensions
                img = Image.open(file_path)
                
                # Create a temporary PDF from the image
                doc = fitz.open()
                page = doc.new_page(width=img.width, height=img.height)
                rect = fitz.Rect(0, 0, img.width, img.height)
                page.insert_image(rect, filename=file_path)
                
                # Use OCRService to extract text
                spans = self.ocr_service.extract_page_spans(page)
                doc.close()
                
                if spans:
                    ocr_text = "\n".join([span.get("text", "").strip() for span in spans if span.get("text", "").strip()])
                    if ocr_text.strip():
                        return ocr_text.strip()
                
            except Exception as fallback_error:
                self.logger.warning(f"OCRService fallback also failed: {fallback_error}")
            
            return f"Error performing OCR on image: {str(e)}"

    def _is_scanned_pdf(self, pdf_path: str) -> bool:
        """Detect if PDF is scanned (image-based) by analyzing text content.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            True if PDF appears to be scanned, False otherwise
        """
        try:
            import fitz
            doc = fitz.open(pdf_path)
            total_chars = 0
            total_pages = len(doc)
            
            # Sample first few pages to determine if scanned
            sample_pages = min(3, total_pages)
            
            for page_num in range(sample_pages):
                page = doc[page_num]
                # Try multiple extraction methods
                page_text = page.get_text() or ""
                page_text += page.get_text("text") or ""
                page_text += page.get_text("raw") or ""
                
                # Also check dict blocks
                try:
                    text_dict = page.get_text("dict", flags=11)
                    for block in text_dict.get("blocks", []):
                        if block.get("type") == 0:  # Text block
                            for line in block.get("lines", []):
                                for span in line.get("spans", []):
                                    page_text += span.get("text", "")
                except:
                    pass
                
                total_chars += len(page_text.strip())
            
            doc.close()
            
            # Threshold: if average chars per page is less than 50, likely scanned
            avg_chars_per_page = total_chars / sample_pages if sample_pages > 0 else 0
            is_scanned = avg_chars_per_page < 50
            
            self.logger.info(
                f"Scanned PDF detection: {total_chars} chars from {sample_pages} pages "
                f"(avg: {avg_chars_per_page:.1f} chars/page) - {'SCANNED' if is_scanned else 'TEXT-BASED'}"
            )
            
            return is_scanned
            
        except Exception as e:
            self.logger.warning(f"Scanned PDF detection failed: {e}")
            # Default to not scanned if detection fails
            return False
    
    def _perform_ocr_on_page(self, page, page_num: int) -> str:
        """Perform OCR on a single page using Tesseract via OCRService.
        
        Args:
            page: PyMuPDF page object
            page_num: Page number (0-indexed)
            
        Returns:
            Extracted text from OCR
        """
        try:
            # Use OCRService's enhanced OCR method
            spans = self.ocr_service.extract_page_spans(page)
            
            if not spans:
                return ""
            
            # Extract text from spans and combine
            ocr_text_parts = []
            for span in spans:
                text = span.get("text", "").strip()
                if text:
                    ocr_text_parts.append(text)
            
            ocr_text = "\n".join(ocr_text_parts)
            
            if ocr_text.strip():
                self.logger.info(f"OCR extracted {len(ocr_text)} characters from page {page_num + 1}")
                return ocr_text
            else:
                return ""
                
        except ImportError:
            self.logger.debug("pytesseract not available for OCR")
            return ""
        except Exception as ocr_err:
            self.logger.warning(f"OCR failed for page {page_num + 1}: {ocr_err}")
            return ""

    @log_method_entry_exit
    @handle_general_operations(severity=ExceptionSeverity.MEDIUM)
    async def extract_all_text_from_pdf_comprehensive(self, pdf_path: str) -> str:
        """Comprehensive text extraction for multiple file types using Factory Pattern.
        
        This method uses the DocumentExtractorFactory to route to appropriate extractors:
        - PDFs: Regular PDFs, Canva PDFs, scanned PDFs (with automatic OCR detection)
          - ALL EXISTING PDF LOGIC IS PRESERVED in PDFExtractor
        - DOCX: Word documents (text and tables)
        - Images: JPEG, PNG, JPG (using OCR)
        
        Args:
            pdf_path: Path to the file (PDF, DOCX, JPEG, PNG, or JPG)
            
        Returns:
            Complete extracted text as a string
        """
        # Use factory to extract text - preserves all existing PDF logic
        return await document_extractor_factory.extract_text(pdf_path)
    
    async def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file (existing comprehensive PDF extraction logic).
        
        This method uses multiple extraction strategies to handle:
        - Regular PDFs with direct text
        - Canva PDFs with XObject Forms
        - PDFs with structured content trees
        - Scanned/image-based PDFs (with automatic OCR detection and processing)
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Complete extracted text as a string
        """
        all_text = ""
        
        try:
            import fitz
            doc = fitz.open(pdf_path)
            
            # Step 1: Detect if PDF is scanned
            is_scanned = self._is_scanned_pdf(pdf_path)
            
            if is_scanned:
                self.logger.info("Detected scanned PDF - using OCR for all pages")
                # For scanned PDFs, use OCR directly for all pages
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    ocr_text = self._perform_ocr_on_page(page, page_num)
                    if ocr_text.strip():
                        all_text += f"\n--- PAGE {page_num + 1} ---\n{ocr_text}\n"
                doc.close()
                
                # Clean and return OCR results
                if all_text.strip():
                    import re
                    cleaned_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', all_text)
                    cleaned_text = re.sub(r'[ \t]+', ' ', cleaned_text)
                    return cleaned_text.strip()
                else:
                    return "No text content found in scanned PDF"
            
            # Step 2: For text-based PDFs, use comprehensive extraction strategies
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_texts = []
                
                # Strategy 1: Standard PyMuPDF extraction methods
                try:
                    page_texts.append(page.get_text())
                    page_texts.append(page.get_text("text"))
                    page_texts.append(page.get_text("raw"))
                    page_texts.append(page.get_text("layout"))
                except Exception as e:
                    self.logger.warning(f"Standard extraction failed for page {page_num + 1}: {e}")
                
                # Strategy 2: Extract from dict blocks (catches structured content)
                try:
                    td = page.get_text("dict", flags=11)  # flags=11 for comprehensive extraction
                    block_lines = []
                    for block in td.get("blocks", []):
                        if block.get("type") == 0:  # Text block
                            for line in block.get("lines", []):
                                line_text = " ".join([span.get("text", "") for span in line.get("spans", [])]).strip()
                                if line_text:
                                    block_lines.append(line_text)
                    if block_lines:
                        page_texts.append("\n".join(block_lines))
                except Exception as e:
                    self.logger.warning(f"Dict block extraction failed for page {page_num + 1}: {e}")
                
                # Strategy 3: Extract from structured content tree (for Canva PDFs with StructTreeRoot)
                try:
                    struct_text = []
                    # Get structured text with comprehensive flags
                    text_dict = page.get_text("dict", flags=11)
                    for block in text_dict.get("blocks", []):
                        if block.get("type") == 0:  # Text block
                            for line in block.get("lines", []):
                                line_parts = []
                                for span in line.get("spans", []):
                                    span_text = span.get("text", "").strip()
                                    if span_text:
                                        line_parts.append(span_text)
                                if line_parts:
                                    struct_text.append(" ".join(line_parts))
                    
                    if struct_text:
                        page_texts.append("\n".join(struct_text))
                except Exception as e:
                    self.logger.warning(f"Structured content extraction failed for page {page_num + 1}: {e}")
                
                # Strategy 4: Extract from XObject Forms (critical for Canva PDFs)
                try:
                    # Canva PDFs store content in XObject Forms referenced in page resources
                    xobject_texts = []
                    
                    # Get page resources to find XObjects
                    try:
                        # Try to extract text with different extraction flags that handle XObjects
                        xobj_text = page.get_text(flags=11)  # Comprehensive flags
                        if xobj_text and xobj_text.strip():
                            xobject_texts.append(xobj_text)
                    except:
                        pass
                    
                    # Also try extracting from the document's structure tree if available
                    try:
                        if doc.is_pdf:
                            # Try to get text from all content streams
                            for xref in page.get_contents():
                                try:
                                    # PyMuPDF should handle XObjects automatically, but we can also
                                    # try to get additional text from content streams
                                    stream_text = page.get_text(flags=11)
                                    if stream_text:
                                        xobject_texts.append(stream_text)
                                except:
                                    pass
                    except:
                        pass
                    
                    if xobject_texts:
                        # Deduplicate and combine
                        unique_xobj = list(set([t.strip() for t in xobject_texts if t.strip()]))
                        if unique_xobj:
                            page_texts.append("\n".join(unique_xobj))
                except Exception as e:
                    self.logger.warning(f"XObject extraction failed for page {page_num + 1}: {e}")
                
                # Combine all extraction results for this page
                combined = "\n".join([t for t in page_texts if t and t.strip()])
                
                # If we got substantial text, add it
                if combined.strip() and len(combined.strip()) > 20:
                    all_text += f"\n--- PAGE {page_num + 1} ---\n{combined}\n"
                # Otherwise try OCR as fallback for individual pages
                elif len(combined.strip()) < 20:
                    self.logger.info(f"Minimal text found on page {page_num + 1}, attempting OCR fallback")
                    ocr_text = self._perform_ocr_on_page(page, page_num)
                    if ocr_text.strip():
                        all_text += f"\n--- PAGE {page_num + 1} ---\n{ocr_text}\n"
            
            doc.close()
            
        except Exception as e:
            self.logger.error(f"PyMuPDF comprehensive extraction failed: {e}")
        
        # Fallback: Try pdfplumber if PyMuPDF didn't get enough
        if len(all_text.strip()) < 200:
            try:
                import pdfplumber
                with pdfplumber.open(pdf_path) as pdf:
                    plumber_text = ""
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            plumber_text += f"\n--- PAGE {page_num + 1} ---\n{page_text}\n"
                    
                    if len(plumber_text.strip()) > len(all_text.strip()):
                        all_text = plumber_text
                        self.logger.info("pdfplumber provided better extraction results")
            except Exception as e:
                self.logger.warning(f"pdfplumber fallback failed: {e}")
        
        # Final fallback: PyPDF2
        if len(all_text.strip()) < 100:
            try:
                import PyPDF2
                with open(pdf_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    pypdf2_text = ""
                    for page_num, page in enumerate(reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            pypdf2_text += f"\n--- PAGE {page_num + 1} ---\n{page_text}\n"
                    
                    if len(pypdf2_text.strip()) > len(all_text.strip()):
                        all_text = pypdf2_text
                        self.logger.info("PyPDF2 provided fallback extraction results")
            except Exception as e:
                self.logger.warning(f"PyPDF2 fallback failed: {e}")
        
        # Clean and return
        if all_text.strip():
            import re
            cleaned_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', all_text)
            cleaned_text = re.sub(r'[ \t]+', ' ', cleaned_text)
            return cleaned_text.strip()
        else:
            return "No text content found in PDF"

    @log_method_entry_exit
    @handle_general_operations(severity=ExceptionSeverity.MEDIUM)
    async def extract_all_text_from_pdf(self, pdf_path: str) -> Dict[str, any]:
        """Extract ALL possible text from PDF using multiple extraction methods.
        
        This method uses up to 6 different extraction techniques to ensure maximum
        text recovery from any type of PDF, including scanned documents, images, and forms.
        
        Args:
            pdf_path: Path to the PDF file to extract text from
            
        Returns:
            Dictionary containing:
                - full_text: Complete extracted text (primary output)
                - extraction_stats: Statistics about the extraction process
                - pages_processed: Number of pages successfully processed
                - total_pages: Total pages in the document
                - character_count: Total characters extracted
                - word_count: Approximate word count
                - error: Any errors that occurred (if any)
        """
        doc = None
        extraction_stats = {
            'total_pages': 0,
            'pages_processed': 0,
            'extraction_methods_used': [],
            'characters_extracted': 0,
            'average_chars_per_page': 0,
            'extraction_errors': []
        }
        
        try:
            # Try multiple PDF loading methods
            try:
                doc = self.ocr_service.open_pdf(pdf_path)
                if not doc:
                    raise ValueError("Failed to open PDF using primary method")
            except Exception as e:
                self.logger.warning(f"Primary PDF loading failed, trying alternative method: {e}")
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(pdf_path)
                    self.logger.info("Successfully opened PDF using PyMuPDF")
                except Exception as e2:
                    self.logger.error(f"All PDF loading methods failed: {e2}")
                    return {
                        "full_text": "", 
                        "error": f"Failed to open PDF: {str(e2)}",
                        "extraction_stats": extraction_stats
                    }
            
            total_pages = len(doc)
            extraction_stats['total_pages'] = total_pages
            self.logger.info(f"Processing {total_pages} pages for complete text extraction")
            
            all_pages_text = []
            
            for page_idx in range(total_pages):
                try:
                    page = doc[page_idx]
                    self.logger.info(f"Extracting text from page {page_idx + 1}/{total_pages}")
                    
                    # Method 1: Native text extraction
                    native_text = page.get_text("text")
                    
                    # Method 2: Text extraction with blocks
                    block_text = self._extract_text_using_blocks(page)
                    
                    # Method 3: Extract spans (catches OCR and all text elements)
                    spans_text = self._extract_text_using_spans(page)
                    
                    # Method 4: Raw text extraction with different parameters
                    raw_text = page.get_text("raw")
                    
                    # Method 5: Text extraction with layout preservation
                    layout_text = page.get_text("layout")
                    
                    # Method 6: Try OCR if available and text is minimal
                    ocr_text = ""
                    if len(''.join([native_text or "", block_text or "", spans_text or "", raw_text or "", layout_text or ""])) < 100:
                        try:
                            ocr_text = self.ocr_service.extract_text_from_image(pdf_path, page_num=page_idx)
                        except:
                            pass
                    
                    # Combine all extraction results
                    combined_text = self._combine_extraction_results([
                        native_text, 
                        block_text, 
                        spans_text,
                        raw_text,
                        layout_text,
                        ocr_text
                    ])
                    
                    if combined_text and combined_text.strip():
                        all_pages_text.append(combined_text.strip())
                        extraction_stats['pages_processed'] += 1
                        self.logger.info(f"Page {page_idx + 1}: Extracted {len(combined_text)} characters")
                    else:
                        self.logger.warning(f"No text found on page {page_idx + 1}")
                        
                except Exception as page_error:
                    error_msg = f"Error processing page {page_idx + 1}: {str(page_error)}"
                    self.logger.error(error_msg, exc_info=True)
                    extraction_stats['extraction_errors'].append(error_msg)
                    
                    # Last resort: try to get any text from the page
                    try:
                        fallback_text = page.get_text("text") or ""
                        if not fallback_text:
                            fallback_text = str(page.get_text("dict"))
                        if fallback_text:
                            all_pages_text.append(f"[PAGE {page_idx + 1} - FALLBACK EXTRACTION]\n{fallback_text}")
                            extraction_stats['pages_processed'] += 1
                    except Exception as fallback_error:
                        self.logger.error(f"Fallback extraction failed: {fallback_error}")
                        all_pages_text.append(f"[PAGE {page_idx + 1} - EXTRACTION FAILED]")
            
            # Combine all pages with clear separation
            full_document_text = "\n\n".join(all_pages_text)
            
            # Final cleanup and deduplication
            full_document_text = self._clean_combined_text(full_document_text)
            
            # Update statistics
            char_count = len(full_document_text)
            extraction_stats.update({
                'characters_extracted': char_count,
                'word_count': len(full_document_text.split()),
                'average_chars_per_page': char_count / max(1, extraction_stats['pages_processed']),
                'extraction_methods_used': ['native', 'blocks', 'spans', 'raw', 'layout', 'ocr']
            })
            
            self.logger.info(
                f"Completed extraction: {char_count} characters "
                f"from {extraction_stats['pages_processed']}/{total_pages} pages"
            )
            
            return {
                "full_text": full_document_text,
                "extraction_stats": extraction_stats,
                "pages_processed": extraction_stats['pages_processed'],
                "total_pages": total_pages,
                "character_count": char_count,
                "word_count": extraction_stats['word_count']
            }
            
            # Combine all pages with clear separation
            full_document_text = "\n\n".join(all_pages_text)
            
            # Final cleanup of the combined text
            full_document_text = self._clean_combined_text(full_document_text)
            
            self.logger.info(
                f"Completed extraction: {len(full_document_text)} characters "
                f"from {len(all_pages_text)}/{total_pages} pages"
            )
            
            return {
                "full_text": full_document_text,
                "pages_processed": len(all_pages_text),
                "total_pages": total_pages,
                "character_count": len(full_document_text),
                "word_count": len(full_document_text.split())
            }
            
        except Exception as e:
            self.logger.error(f"PDF text extraction failed: {e}", exc_info=True)
            return {
                "full_text": "",
                "error": f"Text extraction failed: {str(e)}",
                "pages_processed": len(all_pages_text) if 'all_pages_text' in locals() else 0
            }
            
        finally:
            if doc:
                try:
                    doc.close()
                except:
                    pass
    
    def _extract_text_using_blocks(self, page) -> str:
        """Extract text using block-based method with error handling"""
        try:
            text_dict = page.get_text("dict", flags=0)
            block_text = []
            
            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    block_lines = []
                    for line in block.get("lines", []):
                        line_text = " ".join([span.get("text", "").strip() 
                                           for span in line.get("spans", []) 
                                           if span.get("text", "").strip()])
                        if line_text:
                            block_lines.append(line_text)
                    
                    if block_lines:
                        block_text.append(" ".join(block_lines))
            
            return "\n".join(block_text).strip()
            
        except Exception as e:
            self.logger.warning(f"Block text extraction failed: {e}")
            return ""
    
    def _extract_text_using_spans(self, page) -> str:
        """Extract and format text using spans with enhanced error handling"""
        try:
            spans = self.ocr_service.extract_page_spans(page)
            if not spans:
                return ""
                
            # Sort spans by position
            spans.sort(key=lambda s: (s.get("y0", 0), s.get("x0", 0)))
            
            # Group into lines and then format
            lines = []
            current_line = []
            current_y = spans[0].get("y0", 0) if spans else 0
            
            # Calculate dynamic line tolerance
            if spans:
                heights = [s.get("y1", 0) - s.get("y0", 0) for s in spans 
                         if s.get("y1", 0) > s.get("y0", 0)]
                line_tolerance = max(10, (statistics.median(heights) * 0.8)) if heights else 10
            else:
                line_tolerance = 10
            
            for span in spans:
                span_y = span.get("y0", 0)
                span_text = span.get("text", "").strip()
                
                if not span_text:
                    continue
                    
                if abs(span_y - current_y) <= line_tolerance:
                    current_line.append((span.get("x0", 0), span_text))
                else:
                    if current_line:
                        lines.append(" ".join(t[1] for t in sorted(current_line, key=lambda x: x[0])))
                    current_line = [(span.get("x0", 0), span_text)]
                    current_y = span_y
            
            # Add the last line
            if current_line:
                lines.append(" ".join(t[1] for t in sorted(current_line, key=lambda x: x[0])))
            
            return "\n".join(lines).strip()
            
        except Exception as e:
            self.logger.warning(f"Span-based extraction failed: {e}")
            return ""
    
    def _combine_extraction_results(self, texts: List[str]) -> str:
        """Intelligently combine multiple extraction results to maximize text recovery.
        
        This method uses several techniques to ensure no text is lost during combination:
        1. Removes exact duplicates
        2. Preserves unique content from all sources
        3. Handles different text orderings
        4. Maintains paragraph structure
        
        Args:
            texts: List of text strings from different extraction methods
            
        Returns:
            Combined text with all unique content preserved
        """
        if not texts:
            return ""
            
        # Remove empty strings and normalize whitespace
        texts = [' '.join(str(t).split()) for t in texts if t and str(t).strip()]
        
        if not texts:
            return ""
            
        # Remove exact duplicates while preserving order
        seen_texts = set()
        unique_texts = []
        for text in texts:
            if text not in seen_texts:
                seen_texts.add(text)
                unique_texts.append(text)
        texts = unique_texts
        
        if len(texts) == 1:
            return texts[0]
            
        # Start with the longest text as base
        combined = max(texts, key=len)
        
        # For each other text, find and add unique content
        for text in texts:
            if text == combined:
                continue
                
            # If text is significantly different, add its unique paragraphs
            if len(text) > len(combined) * 0.7:  # Only consider substantial texts
                # Split into paragraphs and sentences for more granular comparison
                combined_paras = set(p.strip() for p in combined.split('\n\n') if p.strip())
                text_paras = [p.strip() for p in text.split('\n\n') if p.strip()]
                
                # Add paragraphs that contain substantial new content
                for para in text_paras:
                    if not any(para in cp or cp in para for cp in combined_paras if len(cp) > 20):
                        combined += "\n\n" + para
                        combined_paras.add(para)
                
                # Also check for unique sentences within paragraphs
                combined_sents = set()
                for para in combined_paras:
                    combined_sents.update(s.strip() for s in re.split(r'(?<=[.!?])\s+', para) if len(s.strip()) > 10)
                
                for para in text_paras:
                    for sent in re.split(r'(?<=[.!?])\s+', para):
                        sent = sent.strip()
                        if len(sent) > 20 and not any(sent in cs or cs in sent for cs in combined_sents):
                            combined += " " + sent
                            combined_sents.add(sent)
        
        # Final cleanup
        return re.sub(r'\s+', ' ', combined).strip()
    
    def _clean_combined_text(self, text: str) -> str:
        """Clean, normalize, and optimize the final combined text.
        
        This method performs several cleaning steps:
        1. Normalizes all whitespace and line breaks
        2. Removes duplicate content while preserving order
        3. Fixes common OCR artifacts
        4. Optimizes paragraph structure
        
        Args:
            text: The text to clean
            
        Returns:
            Cleaned and normalized text with optimal formatting
        """
        if not text or not text.strip():
            return ""
            
        # Convert to string in case input is not a string
        text = str(text)
        
        # Normalize all whitespace and line breaks
        text = re.sub(r'[\r\n]+', '\n', text)  # Normalize line endings
        text = re.sub(r'[\t\f\v ]+', ' ', text)  # Normalize spaces
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize paragraph breaks
        
        # Fix common OCR artifacts and formatting issues
        text = self._fix_common_ocr_errors(text)
        
        # Split into paragraphs and process each
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        # Remove duplicate paragraphs while preserving order
        seen_paragraphs = set()
        unique_paragraphs = []
        
        for para in paragraphs:
            # Skip empty or very short paragraphs
            if not para or len(para) < 3:
                continue
                
            # Normalize the paragraph for comparison
            normalized = ' '.join(para.split())
            
            # Skip if we've seen this exact paragraph before
            if normalized in seen_paragraphs:
                continue
                
            # Skip if this paragraph is contained within another paragraph
            if any(normalized in p and len(normalized) < len(p) for p in seen_paragraphs):
                continue
                
            # Remove any paragraphs that are contained within this one
            seen_paragraphs = {p for p in seen_paragraphs if p not in normalized or p == normalized}
            
            seen_paragraphs.add(normalized)
            unique_paragraphs.append(para)
        
        # Join with double newlines to maintain paragraph structure
        result = '\n\n'.join(unique_paragraphs)
        
        # Final whitespace cleanup
        result = re.sub(r'\s+', ' ', result).strip()
        
        return result
    
    def _format_spans_as_plain_text(self, spans):
        """Format OCR spans into readable plain text"""
        if not spans:
            return ""
        
        # Sort spans by position (top to bottom, left to right)
        spans.sort(key=lambda s: (s.get("y0", 0), s.get("x0", 0)))
        
        # Group into lines with adaptive tolerance
        lines = []
        current_line = []
        current_y = spans[0].get("y0", 0) if spans else 0
        
        # Calculate dynamic line tolerance based on span heights
        if spans:
            heights = [abs(s.get("y1", 0) - s.get("y0", 0)) for s in spans if s.get("y1", 0) > s.get("y0", 0)]
            if heights:
                avg_height = sum(heights) / len(heights)
                line_tolerance = max(15, avg_height * 1.5)  # Adaptive tolerance
            else:
                line_tolerance = 15
        else:
            line_tolerance = 15
        
        for span in spans:
            span_y = span.get("y0", 0)
            span_text = span.get("text", "").strip()
            
            if not span_text:
                continue
                
            # Check if span is on the same line
            if abs(span_y - current_y) < line_tolerance:
                current_line.append(span_text)
            else:
                # Finalize current line
                if current_line:
                    lines.append(" ".join(current_line))
                current_line = [span_text]
                current_y = span_y
        
        # Add the last line
        if current_line:
            lines.append(" ".join(current_line))
        
        return "\n".join(lines)

    @log_method_entry_exit
    @handle_general_operations(severity=ExceptionSeverity.MEDIUM)
    def extract_text_from_regions(self, text_regions: List[Dict]) -> Dict[str, any]:
        """Extract and organize text from text regions with comprehensive content"""
        try:
            if not text_regions:
                return {
                    "sections": [],
                    "full_text": "",
                    "section_count": 0,
                    "word_count": 0,
                    "character_count": 0
                }
            
            extracted_sections = []
            full_text_parts = []
            total_words = 0
            total_chars = 0
            
            for region in text_regions:
                text_content = self._extract_text_from_region(region)
                if text_content:
                    extracted_sections.append(text_content)
                    full_text_parts.append(text_content.get('raw_text', ''))
                    total_words += len(text_content.get('cleaned_text', '').split())
                    total_chars += len(text_content.get('cleaned_text', ''))
            
            full_text = '\n\n'.join(filter(None, full_text_parts))
            
            self.logger.info(
                f"Extracted {len(extracted_sections)} text sections with "
                f"{total_words} words and {total_chars} characters in total"
            )
            
            return {
                "sections": extracted_sections,
                "full_text": full_text,
                "section_count": len(extracted_sections),
                "word_count": total_words,
                "character_count": total_chars
            }
            
        except Exception as e:
            self.logger.error(f"Text extraction from regions failed: {e}")
            return []
            
    @log_method_entry_exit
    @handle_general_operations(severity=ExceptionSeverity.MEDIUM)
    def extract_complete_text(self, text_regions: List[Dict]) -> Dict[str, any]:
        """Extract complete formatted text from text regions with comprehensive content handling
        
        Args:
            text_regions: List of text regions to extract text from
            
        Returns:
            Dict containing:
                - formatted_text: Complete concatenated text from all regions
                - word_count: Total word count
                - character_count: Total character count (including spaces)
        """
        try:
            if not text_regions:
                return {
                    "formatted_text": "",
                    "word_count": 0,
                    "character_count": 0
                }
            
            # First, try to extract text using the existing method
            try:
                result = self.extract_text_from_regions(text_regions)
                if result and result.get("full_text"):
                    return {
                        "formatted_text": result["full_text"],
                        "word_count": result.get("word_count", 0),
                        "character_count": result.get("character_count", 0)
                    }
            except Exception as e:
                self.logger.warning(f"Primary text extraction failed, falling back to raw extraction: {e}")
            
            # Fallback: Raw text extraction if the primary method fails
            all_text_parts = []
            
            for region in text_regions:
                if not isinstance(region, dict):
                    continue
                    
                # Try to get spans first
                spans = region.get("spans", [])
                if spans:
                    # Extract text from spans
                    for span in spans:
                        if isinstance(span, dict) and "text" in span:
                            text = str(span["text"]).strip()
                            if text:
                                all_text_parts.append(text)
                else:
                    # Fallback to direct text extraction
                    text = region.get("text")
                    if text:
                        all_text_parts.append(str(text).strip())
            
            # Join all text parts with spaces
            full_text = " ".join(filter(None, all_text_parts))
            
            # Clean up multiple spaces
            full_text = " ".join(full_text.split())
            
            # Calculate metrics
            word_count = len(full_text.split())
            char_count = len(full_text)
            
            self.logger.info(
                f"Extracted complete text with {word_count} words and {char_count} characters"
            )
            
            return {
                "formatted_text": full_text,
                "word_count": word_count,
                "character_count": char_count
            }
            
        except Exception as e:
            self.logger.error(f"Complete text extraction failed: {e}", exc_info=True)
            # Last resort: return raw text from all regions
            try:
                raw_text = " ".join(str(r.get("text", "")).strip() for r in text_regions if r and isinstance(r, dict))
                word_count = len(raw_text.split())
                return {
                    "formatted_text": raw_text,
                    "word_count": word_count,
                    "character_count": len(raw_text)
                }
            except:
                return {
                    "formatted_text": "",
                    "word_count": 0,
                    "character_count": 0
                }

    def _extract_text_from_region(self, region: Dict) -> Optional[Dict]:
        """Extract text from a single region with enhanced content handling"""
        try:
            spans = region.get("spans", [])
            if not spans:
                return None
            
            # Sort spans by reading order (top to bottom, left to right)
            sorted_spans = self._sort_spans_by_reading_order(spans)
            
            # Extract text with better handling of whitespace and line breaks
            raw_text_parts = []
            prev_y0 = None
            
            for i, span in enumerate(sorted_spans):
                text = span.get("text", "").strip()
                if not text:
                    continue
                    
                # Add newline if we've moved to a new line
                current_y0 = span.get("y0", 0)
                if prev_y0 is not None and abs(current_y0 - prev_y0) > 2:  # Small threshold for same line
                    raw_text_parts.append("\n")
                
                # Add space between words on same line if needed
                if raw_text_parts and not raw_text_parts[-1].endswith("\n") and not text.startswith(" ") and not text.startswith("\n"):
                    raw_text_parts.append(" ")
                
                raw_text_parts.append(text)
                prev_y0 = current_y0
            
            raw_text = "".join(raw_text_parts).strip()
            
            # Clean and normalize the text
            cleaned_text = self._clean_extracted_text(raw_text)
            
            # Analyze text structure
            text_structure = self._analyze_text_structure(sorted_spans)
            
            # Determine text type with enhanced detection
            text_type = self._determine_text_type(cleaned_text, text_structure)
            
            # Calculate metrics
            word_count = len(cleaned_text.split())
            char_count = len(cleaned_text)
            
            return {
                "raw_text": raw_text,
                "cleaned_text": cleaned_text,
                "text_type": text_type,
                "structure": text_structure,
                "bounds": region.get("bounds", {}),
                "confidence": region.get("confidence", 0.0),
                "spans": sorted_spans,
                "metrics": {
                    "word_count": word_count,
                    "character_count": char_count,
                    "line_count": raw_text.count('\n') + 1 if raw_text else 0
                },
                "position": {
                    "page": sorted_spans[0].get("page", 0) if sorted_spans else 0,
                    "y0": min(s.get("y0", 0) for s in sorted_spans) if sorted_spans else 0,
                    "y1": max(s.get("y1", 0) for s in sorted_spans) if sorted_spans else 0
                }
            }
            
        except Exception as e:
            self.logger.warning(f"Text extraction from region failed: {e}")
            return None

    def _sort_spans_by_reading_order(self, spans: List[Dict]) -> List[Dict]:
        """Sort spans by reading order (top to bottom, left to right)"""
        try:
            # First sort by y-coordinate (top to bottom)
            # Then by x-coordinate (left to right) for spans on the same line
            return sorted(spans, key=lambda s: (s.get("y0", 0), s.get("x0", 0)))
            
        except Exception as e:
            self.logger.warning(f"Span sorting failed: {e}")
            return spans

    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text with enhanced processing"""
        try:
            if not text:
                return ""
            
            # First, normalize whitespace and line breaks
            text = ' '.join(text.split())  # Replace all whitespace with single spaces
            
            # Handle common OCR artifacts
            text = self._fix_common_ocr_errors(text)
            
            # Normalize quotes, dashes, and other special characters
            text = self._normalize_special_chars(text)
            
            # Use dynamic text processing service if available
            try:
                from app.Services.DynamicTextProcessingService import DynamicTextProcessingService
                dynamic_processor = DynamicTextProcessingService()
                text = dynamic_processor.process_text(text)
            except ImportError:
                self.logger.debug("DynamicTextProcessingService not available, using basic cleaning")
            
            # Final cleanup
            text = text.strip()
            
            # Ensure proper spacing after punctuation
            text = re.sub(r'([.,!?])([^\s])', r'\1 \2', text)
            
            return text
            
        except Exception as e:
            self.logger.warning(f"Text cleaning failed: {e}")
            return text if text else ""
            
    def _fix_common_ocr_errors(self, text: str) -> str:
        """Fix common OCR errors and artifacts"""
        if not text:
            return ""
            
        # Common OCR replacements
        replacements = {
            '|': 'I',  # Common OCR error for capital I
            '\u2018': "'",  # Left single quote
            '\u2019': "'",  # Right single quote
            '\u201c': '"',  # Left double quote
            '\u201d': '"',  # Right double quote
            '\u2013': '-',   # En dash
            '\u2014': '--',  # Em dash
            '\u2026': '...', # Ellipsis
            '\u00a0': ' ',   # Non-breaking space
            '\u200b': '',    # Zero-width space
        }
        
        # Apply replacements
        for old, new in replacements.items():
            text = text.replace(old, new)
            
        # Fix common word errors
        common_errors = {
            r'\b([A-Z])\.\s+([A-Z])\.': r'\1.\2.',  # Fix initials like "A. B. Smith"
            r'\s+': ' ',  # Replace multiple spaces with single space
            r'\s+([.,!?])': r'\1',  # Remove spaces before punctuation
            r'([a-z])([A-Z])': r'\1 \2'  # Add space between lower and uppercase letters
        }
        
        for pattern, replacement in common_errors.items():
            text = re.sub(pattern, replacement, text)
            
        return text
        
    def _normalize_special_chars(self, text: str) -> str:
        """Normalize special characters and quotes"""
        if not text:
            return ""
            
        # Normalize quotes
        text = text.replace('``', '"')
        text = text.replace("''", '"')
        
        # Normalize dashes
        text = re.sub(r'[\u2010-\u2015]', '-', text)  # Replace all dash types with simple dash
        
        # Normalize whitespace characters
        text = re.sub(r'[\t\r\f\v]+', ' ', text)
        
        return text

    def _analyze_text_structure(self, spans: List[Dict]) -> Dict:
        """Analyze the structure of the text"""
        try:
            if not spans:
                return {}
            
            # Calculate line breaks
            lines = self._group_spans_into_lines(spans)
            
            # Analyze paragraph structure
            paragraphs = self._group_lines_into_paragraphs(lines)
            
            # Calculate text metrics
            total_text = " ".join(span.get("text", "") for span in spans)
            
            return {
                "line_count": len(lines),
                "paragraph_count": len(paragraphs),
                "word_count": len(total_text.split()),
                "character_count": len(total_text),
                "avg_line_length": self._calculate_avg_line_length(lines),
                "has_paragraphs": len(paragraphs) > 1,
                "lines": lines,
                "paragraphs": paragraphs
            }
            
        except Exception as e:
            self.logger.warning(f"Text structure analysis failed: {e}")
            return {}

    def _group_spans_into_lines(self, spans: List[Dict]) -> List[List[Dict]]:
        """Group spans into lines based on vertical position"""
        try:
            if not spans:
                return []
            
            # Sort spans by y-coordinate
            sorted_spans = sorted(spans, key=lambda s: s.get("y0", 0))
            
            lines = []
            current_line = [sorted_spans[0]]
            current_y = sorted_spans[0].get("y0", 0)
            
            for span in sorted_spans[1:]:
                span_y = span.get("y0", 0)
                
                # If span is on the same line (within dynamic tolerance)
                tolerance = self._calculate_dynamic_line_tolerance(spans)
                if abs(span_y - current_y) < tolerance:
                    current_line.append(span)
                else:
                    # Sort current line by x-coordinate
                    current_line.sort(key=lambda s: s.get("x0", 0))
                    lines.append(current_line)
                    current_line = [span]
                    current_y = span_y
            
            # Add the last line
            if current_line:
                current_line.sort(key=lambda s: s.get("x0", 0))
                lines.append(current_line)
            
            return lines
            
        except Exception as e:
            self.logger.warning(f"Line grouping failed: {e}")
            return []

    def _group_lines_into_paragraphs(self, lines: List[List[Dict]]) -> List[List[List[Dict]]]:
        """Group lines into paragraphs based on spacing"""
        try:
            if not lines:
                return []
            
            paragraphs = []
            current_paragraph = [lines[0]]
            
            for i in range(1, len(lines)):
                current_line = lines[i]
                previous_line = lines[i-1]
                
                # Calculate vertical gap between lines
                current_y = min(span.get("y0", 0) for span in current_line)
                previous_y = max(span.get("y1", 0) for span in previous_line)
                gap = current_y - previous_y
                
                # If gap is large, start new paragraph (dynamic threshold)
                threshold = self._calculate_dynamic_paragraph_threshold(lines)
                if gap > threshold:
                    paragraphs.append(current_paragraph)
                    current_paragraph = [current_line]
                else:
                    current_paragraph.append(current_line)
            
            # Add the last paragraph
            if current_paragraph:
                paragraphs.append(current_paragraph)
            
            return paragraphs
            
        except Exception as e:
            self.logger.warning(f"Paragraph grouping failed: {e}")
            return []

    def _calculate_avg_line_length(self, lines: List[List[Dict]]) -> float:
        """Calculate average line length in characters"""
        try:
            if not lines:
                return 0.0
            
            total_length = 0
            for line in lines:
                line_text = " ".join(span.get("text", "") for span in line)
                total_length += len(line_text)
            
            return total_length / len(lines)
            
        except:
            return 0.0

    def _determine_text_type(self, text: str, structure: Dict) -> str:
        """Determine the type of text content dynamically"""
        try:
            if not text:
                return "empty"
            
            # Use dynamic text processing service for type determination
            from app.Services.DynamicTextProcessingService import DynamicTextProcessingService
            dynamic_processor = DynamicTextProcessingService()
            
            # Analyze text characteristics
            characteristics = dynamic_processor._analyze_text_characteristics(text)
            
            # Determine text type dynamically
            text_type = dynamic_processor._determine_text_type_dynamically(characteristics)
            
            # Adjust based on structure dynamically
            paragraph_count = structure.get("paragraph_count", 0)
            line_count = structure.get("line_count", 0)
            
            # Calculate dynamic thresholds based on document characteristics
            if paragraph_count > self._calculate_dynamic_paragraph_threshold_for_type(structure):
                return "paragraph_text"
            elif line_count > self._calculate_dynamic_line_threshold_for_type(structure):
                return "multi_line_text"
            else:
                return text_type
                
        except Exception as e:
            self.logger.warning(f"Dynamic text type determination failed: {e}")
            return "unknown"

    @log_method_entry_exit
    @handle_general_operations(severity=ExceptionSeverity.MEDIUM)
    def format_extracted_text(self, extracted_texts: List[Dict]) -> Dict:
        """Format extracted text into a structured document"""
        try:
            if not extracted_texts:
                return {
                    "formatted_text": "",
                    "sections": [],
                    "total_words": 0,
                    "total_characters": 0
                }
            
            # Sort texts by position (top to bottom)
            sorted_texts = sorted(extracted_texts, 
                                key=lambda x: (x.get("bounds", {}).get("min_y", 0), 
                                             x.get("bounds", {}).get("min_x", 0)))
            
            # Format each section
            sections = []
            total_words = 0
            total_characters = 0
            
            for i, text_content in enumerate(sorted_texts):
                section = {
                    "section_id": i + 1,
                    "text_type": text_content.get("text_type", "unknown"),
                    "content": text_content.get("cleaned_text", ""),
                    "structure": text_content.get("structure", {}),
                    "bounds": text_content.get("bounds", {}),
                    "confidence": text_content.get("confidence", 0.0)
                }
                
                sections.append(section)
                
                # Update totals
                total_words += section["structure"].get("word_count", 0)
                total_characters += section["structure"].get("character_count", 0)
            
            # Create formatted text
            formatted_text = self._create_formatted_text(sections)
            
            result = {
                "formatted_text": formatted_text,
                "sections": sections,
                "total_words": total_words,
                "total_characters": total_characters,
                "section_count": len(sections)
            }
            
            self.logger.info(f"Formatted text with {len(sections)} sections, {total_words} words")
            return result
            
        except Exception as e:
            self.logger.error(f"Text formatting failed: {e}")
            return {
                "formatted_text": "",
                "sections": [],
                "total_words": 0,
                "total_characters": 0
            }

    def _create_formatted_text(self, sections: List[Dict]) -> str:
        """Create formatted text from sections"""
        try:
            formatted_lines = []
            
            for section in sections:
                content = section.get("content", "")
                text_type = section.get("text_type", "")
                
                if not content:
                    continue
                
                # Add section based on type
                if text_type == "document_section":
                    formatted_lines.append(f"\n{content.upper()}\n")
                elif text_type == "formal_text":
                    formatted_lines.append(f"{content}\n")
                elif text_type == "paragraph_text":
                    formatted_lines.append(f"{content}\n\n")
                else:
                    formatted_lines.append(f"{content}\n")
            
            return "\n".join(formatted_lines)
            
        except Exception as e:
            self.logger.warning(f"Formatted text creation failed: {e}")
            return ""

    def _calculate_dynamic_line_tolerance(self, spans: List[Dict]) -> float:
        """Calculate dynamic line tolerance based on span characteristics"""
        try:
            if not spans:
                return 10.0  # Default fallback
            
            # Calculate average span height
            heights = [s.get("y1", 0) - s.get("y0", 0) for s in spans if s.get("y1", 0) > s.get("y0", 0)]
            if not heights:
                return 10.0
            
            avg_height = sum(heights) / len(heights)
            
            # Dynamic tolerance based on average height
            tolerance = max(5.0, avg_height * 0.3)  # 30% of average height, minimum 5px
            
            return min(tolerance, 25.0)  # Cap at 25px
            
        except Exception as e:
            self.logger.warning(f"Dynamic line tolerance calculation failed: {e}")
            return 10.0

    def _calculate_dynamic_paragraph_threshold(self, lines: List[List[Dict]]) -> float:
        """Calculate dynamic paragraph threshold based on line characteristics"""
        try:
            if not lines or len(lines) < 2:
                return 20.0  # Default fallback
            
            # Calculate gaps between consecutive lines
            gaps = []
            for i in range(1, len(lines)):
                current_line = lines[i]
                previous_line = lines[i-1]
                
                current_y = min(span.get("y0", 0) for span in current_line)
                previous_y = max(span.get("y1", 0) for span in previous_line)
                gap = current_y - previous_y
                
                if gap > 0:
                    gaps.append(gap)
            
            if not gaps:
                return 20.0
            
            # Calculate dynamic threshold based on gap statistics
            avg_gap = sum(gaps) / len(gaps)
            gap_std = (sum((g - avg_gap) ** 2 for g in gaps) / len(gaps)) ** 0.5
            
            # Threshold is average gap plus one standard deviation
            threshold = avg_gap + gap_std
            
            return max(threshold, 15.0)  # Minimum 15px
            
        except Exception as e:
            self.logger.warning(f"Dynamic paragraph threshold calculation failed: {e}")
            return 20.0

    def _calculate_dynamic_paragraph_threshold_for_type(self, structure: Dict) -> int:
        """Calculate dynamic paragraph threshold for text type determination"""
        try:
            # Base threshold
            base_threshold = 3
            
            # Adjust based on document characteristics
            word_count = structure.get("word_count", 0)
            character_count = structure.get("character_count", 0)
            
            # If document is very long, be more lenient
            if word_count > 1000 or character_count > 5000:
                return base_threshold - 1
            elif word_count < 100 or character_count < 500:
                return base_threshold + 1
            else:
                return base_threshold
                
        except Exception as e:
            self.logger.warning(f"Dynamic paragraph threshold for type calculation failed: {e}")
            return 3

    def _calculate_dynamic_line_threshold_for_type(self, structure: Dict) -> int:
        """Calculate dynamic line threshold for text type determination"""
        try:
            # Base threshold
            base_threshold = 5
            
            # Adjust based on document characteristics
            word_count = structure.get("word_count", 0)
            avg_line_length = structure.get("avg_line_length", 0)
            
            # If lines are very long, be more lenient
            if avg_line_length > 100:
                return base_threshold - 1
            elif avg_line_length < 20:
                return base_threshold + 1
            elif word_count < 50:
                return base_threshold + 2
            else:
                return base_threshold
                
        except Exception as e:
            self.logger.warning(f"Dynamic line threshold for type calculation failed: {e}")
            return 5
